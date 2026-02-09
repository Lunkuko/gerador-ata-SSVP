import streamlit as st
import streamlit_authenticator as stauth
from streamlit_gsheets import GSheetsConnection
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from num2words import num2words
from datetime import datetime, date, timedelta, time
import io
import urllib.parse
import time

# ==============================================================================
# 1. CONFIGURAÇÃO E CONEXÃO
# ==============================================================================
st.set_page_config(page_title="Gerador de Ata SSVP (Seguro)", layout="wide", page_icon="✝️")

try:
    conn = st.connection("gsheets", type=GSheetsConnection)
except Exception as e:
    st.error("Erro de conexão. Verifique se o arquivo .streamlit/secrets.toml existe.")
    st.stop()

# ==============================================================================
# 2. SISTEMA DE LOGIN E SEGURANÇA
# ==============================================================================
def carregar_usuarios():
    """Lê os usuários da planilha 'Usuarios'."""
    try:
        df = conn.read(worksheet="Usuarios", ttl=0) # ttl=0 garante leitura fresca
        df = df.dropna(subset=['username']) 
        
        credentials = {"usernames": {}}
        for _, row in df.iterrows():
            credentials["usernames"][row['username']] = {
                "name": row['name'],
                "password": row['password'], 
                "roles": [row['role']] if 'role' in row else ['editor']
            }
        return credentials
    except:
        return {"usernames": {}}

def salvar_novo_usuario(username, name, password_hash, role):
    """Cria um novo usuário na planilha."""
    try:
        df = conn.read(worksheet="Usuarios", ttl=0)
        if not df.empty and username in df['username'].values:
            return False, "Usuário já existe!"
            
        novo_user = pd.DataFrame([{
            "username": username,
            "name": name,
            "password": password_hash,
            "role": role
        }])
        
        df_atualizado = pd.concat([df, novo_user], ignore_index=True)
        conn.update(worksheet="Usuarios", data=df_atualizado)
        return True, "Usuário criado com sucesso!"
    except Exception as e:
        return False, f"Erro ao salvar: {e}"

# --- Configuração do Autenticador ---
credentials_dict = carregar_usuarios()

if not credentials_dict["usernames"]:
    st.warning("⚠️ Nenhuma aba 'Usuarios' encontrada ou está vazia.")
    st.info("Crie a aba 'Usuarios' com colunas: username, name, password, role")

authenticator = stauth.Authenticate(
    credentials_dict,
    "ssvp_cookie_seguro", 
    "chave_secreta_aleatoria_ssvp_2026", 
    30 
)

name, authentication_status, username = authenticator.login("main")

# ==============================================================================
# 3. VERIFICAÇÃO DE ACESSO
# ==============================================================================
if authentication_status == False:
    st.error("Usuário ou senha incorretos")
    
elif authentication_status == None:
    st.warning("Por favor, faça login para acessar o sistema.")

elif authentication_status:
    # ==========================================================================
    # === ÁREA SEGURA ===
    # ==========================================================================
    
    with st.sidebar:
        st.write(f"👤 Olá, **{name}**")
        authenticator.logout("Sair", "sidebar")
        st.divider()
        
        user_role = credentials_dict['usernames'][username].get('roles', ['editor'])[0]
        
        if user_role == 'admin':
            with st.expander("🔐 Gestão de Usuários (Admin)"):
                st.info("Cadastrar novo acesso")
                with st.form("form_novo_user"):
                    new_user = st.text_input("Login")
                    new_name = st.text_input("Nome")
                    new_pass = st.text_input("Senha", type="password")
                    new_role = st.selectbox("Permissão", ["editor", "admin"])
                    btn_criar = st.form_submit_button("Criar Usuário")
                    
                    if btn_criar:
                        if new_user and new_pass:
                            try:
                                hashed = stauth.Hasher([new_pass]).generate()[0]
                            except:
                                hashed = stauth.Hasher().generate([new_pass])[0]
                                
                            ok, msg = salvar_novo_usuario(new_user, new_name, hashed, new_role)
                            if ok:
                                st.success(msg)
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error(msg)

    # ==========================================================================
    # 4. FUNÇÕES DO SISTEMA (GERADOR DE ATA)
    # ==========================================================================
    
    @st.cache_data(ttl=3600)
    def carregar_dados_cloud():
        tentativas = 0
        max_tentativas = 3
        while tentativas < max_tentativas:
            try:
                df_config = conn.read(worksheet="Config")
                df_membros = conn.read(worksheet="Membros")
                df_anos = conn.read(worksheet="Anos")
                break 
            except Exception as e:
                erro_str = str(e)
                if "429" in erro_str or "Quota exceeded" in erro_str:
                    tentativas += 1
                    time.sleep(2 ** tentativas)
                    if tentativas == max_tentativas:
                        st.error("⚠️ Google sobrecarregado. Aguarde 1 min.")
                        st.stop()
                else:
                    st.error(f"Erro técnico: {e}")
                    st.stop()

        if df_membros.empty: lista_membros = []
        else: lista_membros = df_membros['Nome'].dropna().astype(str).tolist()
            
        if df_anos.empty: lista_anos = []
        else: lista_anos = df_anos['Ano'].dropna().astype(str).tolist()

        config_dict = dict(zip(df_config['Chave'], df_config['Valor']))
        try: config_dict['ultima_ata'] = int(config_dict.get('ultima_ata', 0))
        except: config_dict['ultima_ata'] = 0

        return {"config": config_dict, "membros": lista_membros, "anos": lista_anos}

    def obter_saldo_anterior():
        try:
            # ttl=0 garante que pegamos o saldo da ata que acabou de ser salva
            df_hist = conn.read(worksheet="Historico", ttl=0)
            if not df_hist.empty and 'Saldo' in df_hist.columns and len(df_hist) > 0:
                return float(df_hist['Saldo'].iloc[-1])
        except: pass
        return 0.0

    def buscar_ata_para_edicao(num_ata_busca):
        """Busca ata forçando atualização do cache e limpando formatação."""
        try:
            # 1. Lê a planilha sem cache (ttl=0)
            df_hist = conn.read(worksheet="Historico", ttl=0)
            
            # 2. Limpeza de dados (Converte para String e remove .0 se houver)
            # Ex: "1296.0" vira "1296"
            df_hist['Numero'] = df_hist['Numero'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
            
            # 3. Limpa o termo de busca também
            termo_busca = str(num_ata_busca).strip()
            
            # 4. Filtra
            ata = df_hist[df_hist['Numero'] == termo_busca]
            
            if not ata.empty: 
                return ata.iloc[0].to_dict()
            return None
        except Exception as e: 
            st.error(f"Erro na busca: {e}")
            return None

    def limpar_memoria():
        carregar_dados_cloud.clear()
        st.cache_data.clear()

    def atualizar_config_cloud(chave, valor):
        time.sleep(1) 
        df = conn.read(worksheet="Config")
        if chave in df['Chave'].values:
            df.loc[df['Chave'] == chave, 'Valor'] = str(valor)
        else:
            new_row = pd.DataFrame([{'Chave': chave, 'Valor': str(valor)}])
            df = pd.concat([df, new_row], ignore_index=True)
        conn.update(worksheet="Config", data=df)
        limpar_memoria()

    def gerenciar_lista_cloud(aba, coluna, valor, acao="adicionar"):
        time.sleep(1)
        df = conn.read(worksheet=aba)
        if acao == "adicionar":
            if valor not in df[coluna].values:
                new_row = pd.DataFrame([{coluna: valor}])
                df = pd.concat([df, new_row], ignore_index=True)
                conn.update(worksheet=aba, data=df)
        elif acao == "remover":
            df = df[df[coluna] != valor]
            conn.update(worksheet=aba, data=df)
        limpar_memoria()
        return True

    def salvar_historico_cloud(dados):
        try:
            df_hist = conn.read(worksheet="Historico", ttl=0) # Lê dados frescos
            
            # Normaliza a coluna Numero para comparação
            df_hist['Numero'] = df_hist['Numero'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
            num_atual = str(dados['num_ata']).strip()
            
            nova_linha = {
                "Numero": num_atual,
                "Data": dados['data_reuniao'],
                "Presidente": dados['pres_nome'],
                "Secretario": dados['secretario_nome'],
                "Leitura": dados['leitura_fonte'],
                "Presentes": dados['lista_presentes_txt'],
                "Ausencias": dados['ausencias'],
                "Visitantes": dados['lista_visitantes_txt'],
                "Receita": dados['receita'],
                "Despesa": dados['despesa'],
                "Saldo": dados['saldo'],
                "Socioeconomico": dados['socioeconomico'],
                "Noticias": dados['noticias_trabalhos'],
                "Palavra_Franca": dados['palavra_franca']
            }

            if num_atual in df_hist['Numero'].values:
                idx = df_hist.index[df_hist['Numero'] == num_atual].tolist()[0]
                for col, val in nova_linha.items():
                    df_hist.at[idx, col] = val
                df_atualizado = df_hist
                msg_tipo = "atualizada"
            else:
                df_nova = pd.DataFrame([nova_linha])
                df_atualizado = pd.concat([df_hist, df_nova], ignore_index=True)
                msg_tipo = "criada"

            conn.update(worksheet="Historico", data=df_atualizado)
            return True, msg_tipo
        except Exception as e:
            st.error(f"Erro ao salvar: {e}")
            return False, "erro"

    # --- Utilitários ---
    def obter_proxima_data(dia_alvo):
        if dia_alvo is None or dia_alvo == "": return datetime.now().date()
        try: dia_alvo = int(dia_alvo)
        except: return datetime.now().date()
        hoje = datetime.now().date()
        dia_hoje = hoje.weekday()
        if dia_hoje == dia_alvo: return hoje
        dias_add = (dia_alvo - dia_hoje + 7) % 7
        return hoje + timedelta(days=dias_add)

    def get_index_membro(nome, lista):
        try: return lista.index(nome) if nome in lista else 0
        except: return 0

    def formatar_valor_extenso(valor):
        try:
            extenso = num2words(valor, lang='pt_BR', to='currency')
            return f"R$ {valor:,.2f} ({extenso})".replace(",", "X").replace(".", ",").replace("X", ".")
        except: return "R$ 0,00"

    def formatar_data_br(data):
        if isinstance(data, (datetime, date)): return data.strftime('%d/%m/%Y')
        if not data or str(data) == "nan": return ""
        return str(data)

    # --- Geradores Docs ---
    class PDF(FPDF):
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Página {self.page_no()}/{{nb}}', 0, 0, 'C')

    def gerar_docx(dados):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        
        texto = f"Ata nº {dados['num_ata']} da reunião ordinária da Conferência {dados['conf_nome']} da SSVP"
        if dados['data_fundacao']: texto += f", fundada em {dados['data_fundacao']}"
        if dados['data_agregacao']: texto += f", agregada em {dados['data_agregacao']}"
        texto += f", vinculada ao Conselho Particular {dados['cons_particular']}, área do Central de {dados['cons_central']}, realizada às {dados['hora_inicio']} do dia {dados['data_reuniao']} do Ano Temático: {dados['ano_tematico']}, na sala de reuniões {dados['local']}."
        
        texto += f" Louvado seja nosso Senhor Jesus Cristo! A reunião foi iniciada pelo Presidente, {dados['pres_nome']}, com as orações regulamentares da Sociedade de São Vicente de Paulo-SSVP."
        texto += f" A leitura espiritual foi tirada do(a) {dados['leitura_fonte']}, proclamada pelo(a) Cfd/Csc. {dados['leitor_nome']}, sendo refletida por alguns membros."
        texto += f" A ata anterior foi lida e {dados['status_ata_ant']}."
        texto += f" Em seguida foi feita a chamada, com a presença dos Confrades e Consócias: {dados['lista_presentes_txt']}."
        
        if dados['lista_visitantes_txt']: texto += f" Presenças dos visitantes: {dados['lista_visitantes_txt']}."
        
        rec = formatar_valor_extenso(dados['receita'])
        des = formatar_valor_extenso(dados['despesa'])
        dec = formatar_valor_extenso(dados['decima'])
        sal = formatar_valor_extenso(dados['saldo'])
        tes = f"o(a) Tesoureiro(a) {dados['tes_nome']}" if dados['tes_nome'] else "o Tesoureiro"
        texto += f" Movimento do Caixa: em seguida {tes} apresentou o estado do caixa: Receita total: {rec}; Despesa total: {des}; Décima semanal: {dec}; Saldo final: {sal}."
        
        if dados['lista_visitantes_txt']: texto += " Agradecimentos aos visitantes."
        if dados['socioeconomico']: texto += f" Levantamento Socioeconômico: {dados['socioeconomico']}."
        if dados['noticias_trabalhos']: texto += f" Notícias dos trabalhos da semana: {dados['noticias_trabalhos']}."
        if dados['escala_visitas']: texto += f" Novas nomeações (escala de visitas): {dados['escala_visitas']}."
        if dados['palavra_franca']: texto += f" Palavra franca: {dados['palavra_franca']}."
        if dados['expediente']: texto += f" Expediente: {dados['expediente']}."
        if dados['palavra_visitantes']: texto += f" Palavra dos Visitantes: {dados['palavra_visitantes']}."
        
        tes_col = f"o(a) tesoureiro(a) {dados['tes_nome']}" if dados['tes_nome'] else "o tesoureiro"
        texto += f" Coleta Secreta: em seguida {tes_col} fez a coleta secreta, enquanto os demais cantavam {dados['musica_final']}."
        texto += f" Nada mais havendo a tratar, a reunião foi encerrada com as orações finais regulamentares da SSVP e com a oração para Canonização do Beato Frederico Ozanam, às {dados['hora_fim']}."
        texto += f" Para constar, eu, {dados['secretario_nome']}, {dados['secretario_cargo']}, lavrei a presente ata, que dato e assino."
        
        paragrafo = doc.add_paragraph(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        pd = doc.add_paragraph(f"{dados['cidade_estado']}, {dados['data_reuniao']}.")
        pd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("\n\nAssinaturas dos Presentes:")
        for _ in range(30):
            doc.add_paragraph("___________________________________________________________________________________")

        return doc

    def gerar_pdf_nativo(dados):
        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.set_margins(25, 25, 25)
        
        texto = f"Ata nº {dados['num_ata']} da reunião ordinária da Conferência {dados['conf_nome']} da SSVP"
        if dados['data_fundacao']: texto += f", fundada em {dados['data_fundacao']}"
        if dados['data_agregacao']: texto += f", agregada em {dados['data_agregacao']}"
        texto += f", vinculada ao Conselho Particular {dados['cons_particular']}, área do Central de {dados['cons_central']}, realizada às {dados['hora_inicio']} do dia {dados['data_reuniao']} do Ano Temático: {dados['ano_tematico']}, na sala de reuniões {dados['local']}."
        
        texto += f" Louvado seja nosso Senhor Jesus Cristo! A reunião foi iniciada pelo Presidente, {dados['pres_nome']}, com as orações regulamentares da Sociedade de São Vicente de Paulo-SSVP."
        texto += f" A leitura espiritual foi tirada do(a) {dados['leitura_fonte']}, proclamada pelo(a) Cfd/Csc. {dados['leitor_nome']}, sendo refletida por alguns membros."
        texto += f" A ata anterior foi lida e {dados['status_ata_ant']}."
        texto += f" Em seguida foi feita a chamada, com a presença dos Confrades e Consócias: {dados['lista_presentes_txt']}."
        
        if dados['lista_visitantes_txt']: texto += f" Presenças dos visitantes: {dados['lista_visitantes_txt']}."
        
        rec = formatar_valor_extenso(dados['receita'])
        des = formatar_valor_extenso(dados['despesa'])
        dec = formatar_valor_extenso(dados['decima'])
        sal = formatar_valor_extenso(dados['saldo'])
        tes = f"o(a) Tesoureiro(a) {dados['tes_nome']}" if dados['tes_nome'] else "o Tesoureiro"
        texto += f" Movimento do Caixa: em seguida {tes} apresentou o estado do caixa: Receita total: {rec}; Despesa total: {des}; Décima semanal: {dec}; Saldo final: {sal}."
        
        if dados['lista_visitantes_txt']: texto += " Agradecimentos aos visitantes."
        if dados['socioeconomico']: texto += f" Levantamento Socioeconômico: {dados['socioeconomico']}."
        if dados['noticias_trabalhos']: texto += f" Notícias dos trabalhos da semana: {dados['noticias_trabalhos']}."
        if dados['escala_visitas']: texto += f" Novas nomeações (escala de visitas): {dados['escala_visitas']}."
        if dados['palavra_franca']: texto += f" Palavra franca: {dados['palavra_franca']}."
        if dados['expediente']: texto += f" Expediente: {dados['expediente']}."
        if dados['palavra_visitantes']: texto += f" Palavra dos Visitantes: {dados['palavra_visitantes']}."
        
        tes_col = f"o(a) tesoureiro(a) {dados['tes_nome']}" if dados['tes_nome'] else "o tesoureiro"
        texto += f" Coleta Secreta: em seguida {tes_col} fez a coleta secreta, enquanto os demais cantavam {dados['musica_final']}."
        texto += f" Nada mais havendo a tratar, a reunião foi encerrada com as orações finais regulamentares da SSVP e com a oração para Canonização do Beato Frederico Ozanam, às {dados['hora_fim']}."
        texto += f" Para constar, eu, {dados['secretario_nome']}, {dados['secretario_cargo']}, lavrei a presente ata, que dato e assino."
        
        pdf.multi_cell(0, 7, texto, align="J")
        
        pdf.ln(10)
        pdf.cell(0, 10, f"{dados['cidade_estado']}, {dados['data_reuniao']}.", ln=True, align="R")
        
        pdf.ln(10)
        pdf.cell(0, 10, "Assinaturas dos Presentes:", ln=True, align="L")
        for _ in range(30):
            pdf.cell(0, 8, "_______________________________________________________________________", ln=True, align="C")

        return bytes(pdf.output(dest='S'))

    # ==========================================================================
    # 5. INTERFACE PRINCIPAL (UI)
    # ==========================================================================
    db = carregar_dados_cloud()
    
    if 'dados_carregados' not in st.session_state:
        st.session_state.dados_carregados = {}
    dc = st.session_state.dados_carregados

    dia_cfg = db['config'].get('dia_semana_reuniao', None)
    data_pad = obter_proxima_data(dia_cfg)
    hora_pad_str = db['config'].get('horario_padrao', '20:00')
    try: hora_pad = datetime.strptime(hora_pad_str, '%H:%M').time()
    except: hora_pad = time(20, 0)
    loc_pad = db['config'].get('local_padrao', '')
    cid_pad = db['config'].get('cidade_padrao', '')

    with st.sidebar:
        st.header("⚙️ Painel de Controle")
        
        with st.expander("🛠️ Corrigir Ata Anterior", expanded=True):
            st.info("Digite o nº para editar.")
            nb = st.number_input("Nº Ata", min_value=1, step=1, key="find_ata")
            if st.button("Carregar"):
                d_old = buscar_ata_para_edicao(nb)
                if d_old:
                    st.session_state.dados_carregados = d_old
                    st.toast(f"Ata {nb} carregada!")
                    time.sleep(1)
                    st.rerun()
                else: st.error("Não encontrada.")

        with st.expander("👔 Cargos"):
            ip = get_index_membro(db['config'].get('pres_padrao'), db['membros'])
            cp = st.selectbox("Presidente", db['membros'], index=ip)
            st.divider()
            is1 = get_index_membro(db['config'].get('sec_padrao'), db['membros'])
            cs1 = st.selectbox("1º Sec.", db['membros'], index=is1)
            csc1 = st.text_input("Cargo 1", db['config'].get('sec_cargo_padrao',''))
            st.divider()
            is2 = get_index_membro(db['config'].get('sec2_padrao'), db['membros'])
            cs2 = st.selectbox("2º Sec.", db['membros'], index=is2)
            csc2 = st.text_input("Cargo 2", db['config'].get('sec2_cargo_padrao',''))
            st.divider()
            it = get_index_membro(db['config'].get('tes_padrao'), db['membros'])
            ct = st.selectbox("Tesoureiro", db['membros'], index=it)
            
            if st.button("Salvar Cargos"):
                atualizar_config_cloud('pres_padrao', cp)
                atualizar_config_cloud('sec_padrao', cs1)
                atualizar_config_cloud('sec_cargo_padrao', csc1)
                atualizar_config_cloud('sec2_padrao', cs2)
                atualizar_config_cloud('sec2_cargo_padrao', csc2)
                atualizar_config_cloud('tes_padrao', ct)
                st.rerun()

        with st.expander("🏢 Configs"):
            cn = st.text_input("Nome Conf.", db['config'].get('nome_conf',''))
            ch = st.text_input("Hora", hora_pad_str)
            cl = st.text_input("Local", loc_pad)
            cc = st.text_input("Cidade", cid_pad)
            cp = st.text_input("Cons. Part.", db['config'].get('cons_particular',''))
            cce = st.text_input("Cons. Cent.", db['config'].get('cons_central',''))
            dfu = st.text_input("Dt Fund.", db['config'].get('data_fundacao',''))
            dag = st.text_input("Dt Agreg.", db['config'].get('data_agregacao',''))
            if st.button("Salvar Configs"):
                atualizar_config_cloud('nome_conf', cn)
                atualizar_config_cloud('horario_padrao', ch)
                atualizar_config_cloud('local_padrao', cl)
                atualizar_config_cloud('cidade_padrao', cc)
                atualizar_config_cloud('cons_particular', cp)
                atualizar_config_cloud('cons_central', cce)
                atualizar_config_cloud('data_fundacao', dfu)
                atualizar_config_cloud('data_agregacao', dag)
                st.rerun()

        with st.expander("👥 Membros"):
            nm = st.text_input("Novo Membro")
            if st.button("Add"): 
                if gerenciar_lista_cloud("Membros","Nome",nm,"adicionar"): st.rerun()
            rm = st.selectbox("Remover", ["..."]+db['membros'])
            if st.button("Del"):
                if rm!="..." and gerenciar_lista_cloud("Membros","Nome",rm,"remover"): st.rerun()

        with st.expander("📅 Anos"):
            na = st.text_input("Novo Ano")
            if st.button("Add Ano"):
                if gerenciar_lista_cloud("Anos","Ano",na,"adicionar"): st.rerun()
        
        st.divider()
        if st.button("Forçar Atualização"): limpar_memoria(); st.rerun()

    st.title("Gerador de Ata SSVP ✝️")
    st.caption("Conectado ao Arquivo Digital")

    val_num = int(dc.get('Numero', db['config']['ultima_ata'] + 1))
    val_data = data_pad
    if 'Data' in dc: 
        try: val_data = datetime.strptime(dc['Data'], '%d/%m/%Y').date()
        except: pass
    
    c1, c2, c3 = st.columns(3)
    num_ata = c1.number_input("Número", value=val_num, step=1)
    if dc: st.caption(f"✏️ Editando Ata {val_num}")
    
    ia = 0
    if 'Ano' in dc and dc['Ano'] in db['anos']: ia = db['anos'].index(dc['Ano'])
    ano_tem = c2.selectbox("Ano Temático", db['anos'], index=ia)
    dt_reuniao = c3.date_input("Data", val_data, format="DD/MM/YYYY")

    with st.expander(f"📍 Detalhes: {hora_pad_str} - {loc_pad}", expanded=False):
        cx1, cx2, cx3 = st.columns(3)
        hr_ini = cx1.time_input("Início", hora_pad)
        local_r = cx2.text_input("Local", loc_pad)
        cidade_r = cx3.text_input("Cidade", cid_pad)

    st.divider()
    
    st.subheader("Chamada")
    cp1, cp2 = st.columns(2)
    def_pres = []
    if 'Presentes' in dc:
        def_pres = [p.strip() for p in dc['Presentes'].split(',') if p.strip() in db['membros']]
    
    with cp1:
        presentes = st.multiselect("1️⃣ Quem veio?", db['membros'], default=def_pres)
    
    ausentes = [m for m in db['membros'] if m not in presentes]
    motivos = {}
    with cp2:
        justif = st.multiselect("2️⃣ Quem justificou?", ausentes)
    
    if justif:
        st.caption("Motivos:")
        cols = st.columns(3)
        for i, m in enumerate(justif):
            motivos[m] = cols[i%3].text_input(m, placeholder="Motivo...")

    st.divider()

    st.subheader("Tesouraria")
    cf1, cf2, cf3, cf4 = st.columns(4)
    saldo_ant = obter_saldo_anterior()
    st.caption(f"Saldo Anterior: R$ {saldo_ant:.2f}")
    
    v_rec = float(dc.get('Receita', 0.0))
    v_des = float(dc.get('Despesa', 0.0))
    v_dec = float(dc.get('Decima', 0.0))
    
    rec = cf1.number_input("Receita", value=v_rec, step=0.1)
    des = cf2.number_input("Despesa", value=v_des, step=0.1)
    dec = cf3.number_input("Décima", value=v_dec, step=0.1)
    
    saldo_calc = saldo_ant + rec - des - dec
    saldo = cf4.number_input("Saldo Final", value=saldo_calc, disabled=True)
    
    it = get_index_membro(db['config'].get('tes_padrao'), db['membros'])
    tes_nome = cf4.selectbox("Tesoureiro", db['membros'], index=it)
    
    if saldo < 0: st.error("Caixa Negativo!")
    
    st.divider()

    ce1, ce2, ce3 = st.columns(3)
    ipres = get_index_membro(dc.get('Presidente', db['config'].get('pres_padrao')), db['membros'])
    pres_nome = ce1.selectbox("Presidente", db['membros'], index=ipres)
    font_l = ce2.text_input("Fonte Leitura", value=dc.get('Leitura',''))
    leit_nome = ce3.selectbox("Leitor", db['membros'])
    
    st.divider()
    c_ata_opt, c_ata_txt = st.columns([1, 2])
    with c_ata_opt:
        st_ata = st.radio("Ata Anterior", ["Aprovada sem ressalvas", "Aprovada com ressalvas"])
    txt_res = ""
    with c_ata_txt:
        if st_ata == "Aprovada com ressalvas":
            txt_res = st.text_input("Detalhes", placeholder="O que estava errado?")
            
    st.divider()
    visit = st.text_area("Visitantes", value=dc.get('Visitantes',''))
    
    st.divider()
    st.markdown("### Relatórios")
    socio = st.text_area("Socioeconômico", value=dc.get('Socioeconomico',''), height=100)
    notic = st.text_area("Notícias", value=dc.get('Noticias',''), height=100)
    escal = st.text_area("Escala", value=dc.get('Escala',''))
    palav = st.text_area("Palavra Franca", value=dc.get('Palavra_Franca',''))
    exped = st.text_area("Expediente", value=dc.get('Expediente',''))
    
    st.divider()
    ce1, ce2 = st.columns(2)
    p_vis = ce1.text_input("Palavra Visitantes", "")
    mov_ex = ce2.text_input("Mov. Extra", "Coleta regular")
    ce3, ce4 = st.columns(2)
    music = ce3.text_input("Música", "Hino de Ozanam")
    hr_fim = ce4.time_input("Fim")
    
    st.divider()
    st.markdown("##### ✍️ Assinatura")
    qa = st.radio("Secretário Hoje?", ["1º Secretário", "2º Secretário", "Outro"], horizontal=True)
    
    ns1 = db['config'].get('sec_padrao')
    ns2 = db['config'].get('sec2_padrao')
    
    if qa == "1º Secretário":
        idx_s = get_index_membro(ns1, db['membros'])
        cg_fin = "1º Secretário(a)"
    elif qa == "2º Secretário":
        idx_s = get_index_membro(ns2, db['membros'])
        cg_fin = "2º Secretário(a)"
    else:
        idx_s = get_index_membro(dc.get('Secretario',''), db['membros'])
        cg_fin = "Secretário(a) ad hoc"
        
    cs1, cs2 = st.columns([2,1])
    sec_nom = cs1.selectbox("Nome Secretário", db['membros'], index=idx_s)
    cs2.text_input("Cargo", value=cg_fin, disabled=True)
    
    st.divider()
    lbl_btn = "💾 Atualizar Ata" if dc else "💾 Gerar Ata"
    if st.button(lbl_btn, type="primary"):
        ls_aus = []
        if ausentes:
            for m in ausentes:
                mot = motivos.get(m, "").strip()
                if mot: ls_aus.append(f"{m} ({mot})")
                elif m not in motivos: ls_aus.append(m)
                else: ls_aus.append(f"{m} (Justificado)")
        txt_aus = ", ".join(ls_aus) if ls_aus else "Não houve."
        
        st_fin = st_ata
        if st_ata == "Aprovada com ressalvas" and txt_res: st_fin += f": {txt_res}"
        
        dados_ata = {
            'num_ata': str(num_ata),
            'conf_nome': db['config'].get('nome_conf',''),
            'cons_particular': db['config'].get('cons_particular',''),
            'cons_central': db['config'].get('cons_central',''),
            'data_fundacao': formatar_data_br(db['config'].get('data_fundacao','')),
            'data_agregacao': formatar_data_br(db['config'].get('data_agregacao','')),
            'ano_tematico': ano_tem,
            'data_reuniao': formatar_data_br(dt_reuniao),
            'hora_inicio': hr_ini.strftime('%H:%M'),
            'local': local_r, 'pres_nome': pres_nome,
            'leitura_fonte': font_l, 'leitor_nome': leit_nome,
            'status_ata_ant': st_fin,
            'lista_presentes_txt': ", ".join(presentes),
            'ausencias': txt_aus,
            'lista_visitantes_txt': visit.replace("\n", ", ") if visit else "",
            'receita': rec, 'despesa': des, 'decima': dec, 'saldo': saldo,
            'tes_nome': tes_nome, 'socioeconomico': socio,
            'noticias_trabalhos': notic, 'escala_visitas': escal,
            'palavra_franca': palav, 'expediente': exped,
            'palavra_visitantes': p_vis, 'mov_financeiro_extra': mov_ex,
            'musica_final': music, 'hora_fim': hr_fim.strftime('%H:%M'),
            'secretario_nome': sec_nom, 'secretario_cargo': cg_fin,
            'cidade_estado': cidade_r
        }
        
        with st.spinner("Processando..."):
            ok, tipo = salvar_historico_cloud(dados_ata)
            if ok:
                st.toast(f"✅ Ata {num_ata} {tipo}!")
                if tipo=="criada" and int(num_ata) > db['config']['ultima_ata']:
                    atualizar_config_cloud('ultima_ata', int(num_ata))
                st.session_state.dados_carregados = {}
        
        doc = gerar_docx(dados_ata)
        bio = io.BytesIO()
        doc.save(bio)
        pdf = gerar_pdf_nativo(dados_ata)
        
        st.success(f"Sucesso! Saldo Final: R$ {saldo:.2f}")
        c_down1, c_down2 = st.columns(2)
        c_down1.download_button("📄 Baixar PDF", pdf, f"Ata_{num_ata}.pdf", "application/pdf", type="primary")
        c_down2.download_button("📝 Baixar Word", bio.getvalue(), f"Ata_{num_ata}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")